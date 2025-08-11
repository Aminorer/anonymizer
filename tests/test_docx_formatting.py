from io import BytesIO
import pathlib
import sys

sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

from docx import Document

from backend.anonymizer import RegexAnonymizer


def create_formatted_doc() -> bytes:
    doc = Document()
    section = doc.sections[0]
    hdr = section.header.paragraphs[0]
    hdr.add_run("Email: ")
    bold_run = hdr.add_run("test@example.com")
    bold_run.bold = True

    table = doc.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.add_run("Phone ")
    ital_run = cell_para.add_run("0123456789")
    ital_run.italic = True

    ftr = section.footer.paragraphs[0]
    ftr.add_run("SIREN ")
    under_run = ftr.add_run("123456789")
    under_run.underline = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_formatting_preserved_header_table_footer():
    data = create_formatted_doc()
    anonymizer = RegexAnonymizer()
    anon_bytes, entities, mapping, text = anonymizer.anonymize_docx(data)
    doc = Document(BytesIO(anon_bytes))

    hdr_run = doc.sections[0].header.paragraphs[0].runs[1]
    tbl_run = doc.tables[0].cell(0, 0).paragraphs[0].runs[1]
    ftr_run = doc.sections[0].footer.paragraphs[0].runs[1]

    assert hdr_run.text == "[EMAIL]"
    assert hdr_run.bold
    assert tbl_run.text == "[PHONE]"
    assert tbl_run.italic
    assert ftr_run.text == "[SIREN]"
    assert ftr_run.underline
