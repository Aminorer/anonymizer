from io import BytesIO
import pathlib
import sys

sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from backend.anonymizer import RegexAnonymizer


def create_sample_doc() -> bytes:
    doc = Document()
    doc.core_properties.title = "Titre"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header info"
    section.footer.paragraphs[0].text = "Footer info"

    doc.add_paragraph("Contact: test@example.com")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Phone 0123456789"

    img = Image.new("RGB", (1, 1), color="red")
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    doc.add_picture(buf)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def test_docx_structure_preserved(tmp_path):
    data = create_sample_doc()
    anonymizer = RegexAnonymizer()
    anonymized, entities, mapping, text = anonymizer.anonymize_docx(data)
    orig_doc = Document(BytesIO(data))
    new_doc = Document(BytesIO(anonymized))

    assert orig_doc.core_properties.title == new_doc.core_properties.title
    assert len(orig_doc.paragraphs) == len(new_doc.paragraphs)
    assert len(orig_doc.tables) == len(new_doc.tables)
    orig_imgs = [r for r in orig_doc.part.rels.values() if r.reltype == RT.IMAGE]
    new_imgs = [r for r in new_doc.part.rels.values() if r.reltype == RT.IMAGE]
    assert len(orig_imgs) == len(new_imgs)
    assert (
        len(orig_doc.sections[0].header.paragraphs)
        == len(new_doc.sections[0].header.paragraphs)
    )
    assert (
        len(orig_doc.sections[0].footer.paragraphs)
        == len(new_doc.sections[0].footer.paragraphs)
    )
