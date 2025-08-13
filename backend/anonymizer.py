from dataclasses import dataclass, asdict
import re
from typing import List, Tuple, Optional, Dict, Any
from io import BytesIO
import tempfile
from pathlib import Path
import logging

import pdfplumber
from pdf2docx import parse as pdf2docx_parse
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

@dataclass
class Entity:
    type: str
    value: str
    start: int
    end: int
    page: Optional[int] = None
    x: Optional[float] = None
    y: Optional[float] = None
    width: Optional[float] = None
    height: Optional[float] = None


@dataclass
class RunInfo:
    start: int
    end: int
    page: int
    section: int
    path: Tuple

    def get_run(self, doc: Document):
        try:
            kind = self.path[0]
            if kind == "body":
                _, p_idx, r_idx = self.path
                if p_idx < len(doc.paragraphs) and r_idx < len(doc.paragraphs[p_idx].runs):
                    return doc.paragraphs[p_idx].runs[r_idx]
            elif kind == "table":
                _, t_idx, row_idx, cell_idx, p_idx, r_idx = self.path
                if (t_idx < len(doc.tables) and 
                    row_idx < len(doc.tables[t_idx].rows) and
                    cell_idx < len(doc.tables[t_idx].rows[row_idx].cells)):
                    cell = doc.tables[t_idx].rows[row_idx].cells[cell_idx]
                    if p_idx < len(cell.paragraphs) and r_idx < len(cell.paragraphs[p_idx].runs):
                        return cell.paragraphs[p_idx].runs[r_idx]
            elif kind == "header":
                _, s_idx, p_idx, r_idx = self.path
                if (s_idx < len(doc.sections) and 
                    p_idx < len(doc.sections[s_idx].header.paragraphs) and
                    r_idx < len(doc.sections[s_idx].header.paragraphs[p_idx].runs)):
                    return doc.sections[s_idx].header.paragraphs[p_idx].runs[r_idx]
            elif kind == "footer":
                _, s_idx, p_idx, r_idx = self.path
                if (s_idx < len(doc.sections) and 
                    p_idx < len(doc.sections[s_idx].footer.paragraphs) and
                    r_idx < len(doc.sections[s_idx].footer.paragraphs[p_idx].runs)):
                    return doc.sections[s_idx].footer.paragraphs[p_idx].runs[r_idx]
            
            logger.warning(f"Impossible de récupérer le run pour le path: {self.path}")
            return None
        except (IndexError, AttributeError) as e:
            logger.warning(f"Erreur lors de la récupération du run: {e}")
            return None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert RunInfo to JSON-serializable dict."""
        try:
            return {
                'start': self.start,
                'end': self.end,
                'page': self.page,
                'section': self.section,
                'path': list(self.path)  # Convert tuple to list for JSON
            }
        except Exception as e:
            logger.error(f"Erreur lors de la conversion RunInfo en dict: {e}")
            return {
                'start': self.start,
                'end': self.end,
                'page': self.page,
                'section': self.section,
                'path': []
            }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'RunInfo':
        """Create RunInfo from dict."""
        try:
            return cls(
                start=data.get('start', 0),
                end=data.get('end', 0),
                page=data.get('page', 0),
                section=data.get('section', 0),
                path=tuple(data.get('path', []))  # Convert list back to tuple
            )
        except Exception as e:
            logger.error(f"Erreur lors de la création RunInfo depuis dict: {e}")
            return cls(start=0, end=0, page=0, section=0, path=())


class RegexAnonymizer:
    """Regex based anonymizer for DOCX and PDF files.

    Besides anonymization it also exposes utilities to map positions
    between the plain extracted text and DOCX runs so that callers can
    later locate and modify text while preserving formatting.
    """

    PATTERNS = {
        "EMAIL": re.compile(r"[\w\.-]+@[\w\.-]+", re.IGNORECASE),
        "PHONE": re.compile(r"(?:\+\d{1,3} ?)?\d{10}"),
        "DATE": re.compile(r"\b\d{2}/\d{2}/\d{4}\b"),
        "IBAN": re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b"),
        "SIREN": re.compile(r"\b\d{9}\b"),
        "SIRET": re.compile(r"\b\d{14}\b"),
        # Basic French address: number + street + zip + city
        "ADDRESS": re.compile(r"\d+\s+[\w\s]+,?\s*\d{5}\s+[\w\s]+"),
        # Location placeholder: capitalized words
        "LOC": re.compile(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b"),
    }

    def detect(self, text: str) -> List[Entity]:
        """Detect entities in ``text`` and return their positions."""
        entities: List[Entity] = []
        try:
            for etype, pattern in self.PATTERNS.items():
                for match in pattern.finditer(text):
                    entities.append(
                        Entity(
                            type=etype,
                            value=match.group(),
                            start=match.start(),
                            end=match.end(),
                        )
                    )
        except Exception as e:
            logger.error(f"Erreur lors de la détection d'entités: {e}")
        return entities

    # ------------------------------------------------------------------
    # DOCX utilities
    def docx_text_mapping(self, doc: Document) -> Tuple[str, List[RunInfo]]:
        """Return full text of ``doc`` and mapping to runs with page/section."""

        parts: List[str] = []
        mapping: List[RunInfo] = []
        pos = 0
        page = 0
        section = 0

        def add_run(run, path, page_val, section_val):
            nonlocal pos
            try:
                text = run.text or ""
                start = pos
                end = start + len(text)
                parts.append(text)
                mapping.append(RunInfo(start, end, page_val, section_val, path))
                pos = end
                # Check for page breaks
                if hasattr(run, '_element') and run._element.xpath(".//w:br[@w:type='page']"):
                    return page_val + 1
                return page_val
            except Exception as e:
                logger.warning(f"Erreur lors de l'ajout du run: {e}")
                return page_val

        def add_sep(text):
            nonlocal pos
            parts.append(text)
            pos += len(text)

        try:
            # Process paragraphs
            for p_idx, para in enumerate(doc.paragraphs):
                try:
                    for r_idx, run in enumerate(para.runs):
                        page = add_run(run, ("body", p_idx, r_idx), page, section)
                    if p_idx < len(doc.paragraphs) - 1:
                        add_sep("\n")
                    # Check for section breaks
                    if hasattr(para, '_p') and para._p.xpath("w:pPr/w:sectPr"):
                        section += 1
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement du paragraphe {p_idx}: {e}")
                    continue
            
            if doc.paragraphs:
                add_sep("\n")

            # Process tables
            for t_idx, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for p_idx, para in enumerate(cell.paragraphs):
                                for r_idx, run in enumerate(para.runs):
                                    page = add_run(
                                        run,
                                        ("table", t_idx, row_idx, cell_idx, p_idx, r_idx),
                                        page,
                                        section,
                                    )
                                if p_idx < len(cell.paragraphs) - 1:
                                    add_sep("\n")
                            if cell_idx < len(row.cells) - 1:
                                add_sep("\t")
                        if row_idx < len(table.rows) - 1:
                            add_sep("\n")
                    if t_idx < len(doc.tables) - 1:
                        add_sep("\n")
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement de la table {t_idx}: {e}")
                    continue

            # Process headers
            for sec_idx, sec in enumerate(doc.sections):
                try:
                    if hasattr(sec, 'header') and sec.header:
                        for p_idx, para in enumerate(sec.header.paragraphs):
                            for r_idx, run in enumerate(para.runs):
                                page = add_run(run, ("header", sec_idx, p_idx, r_idx), page, sec_idx)
                            if p_idx < len(sec.header.paragraphs) - 1:
                                add_sep("\n")
                        if sec.header.paragraphs:
                            add_sep("\n")
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement de l'en-tête de section {sec_idx}: {e}")

            # Process footers
            for sec_idx, sec in enumerate(doc.sections):
                try:
                    if hasattr(sec, 'footer') and sec.footer:
                        for p_idx, para in enumerate(sec.footer.paragraphs):
                            for r_idx, run in enumerate(para.runs):
                                page = add_run(run, ("footer", sec_idx, p_idx, r_idx), page, sec_idx)
                            if p_idx < len(sec.footer.paragraphs) - 1:
                                add_sep("\n")
                        if sec.footer.paragraphs:
                            add_sep("\n")
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement du pied de page de section {sec_idx}: {e}")

        except Exception as e:
            logger.error(f"Erreur lors du mapping du texte DOCX: {e}")

        return "".join(parts), mapping

    def _replace_using_mapping(
        self, doc: Document, entities: List[Entity], mapping: List[RunInfo]
    ) -> None:
        try:
            def _original_text(ent: Entity) -> str:
                parts: List[str] = []
                for m in mapping:
                    if m.end <= ent.start or m.start >= ent.end:
                        continue
                    run = m.get_run(doc)
                    if run is None:
                        continue
                    rs = max(ent.start, m.start) - m.start
                    re = min(ent.end, m.end) - m.start
                    try:
                        run_text = run.text or ""
                        if rs < len(run_text) and re <= len(run_text):
                            parts.append(run_text[rs:re])
                    except Exception as e:
                        logger.warning(f"Erreur lors de l'extraction du texte original: {e}")
                        continue
                return "".join(parts)

            for ent in entities:
                try:
                    original = _original_text(ent)
                    replacement = ent.value
                    if replacement == original:
                        replacement = f"[{ent.type}]"

                    first = True
                    for m in mapping:
                        if m.end <= ent.start or m.start >= ent.end:
                            continue
                        run = m.get_run(doc)
                        if run is None:
                            continue
                        rs = max(ent.start, m.start) - m.start
                        re = min(ent.end, m.end) - m.start
                        try:
                            run_text = run.text or ""
                            if first:
                                run.text = run_text[:rs] + replacement + run_text[re:]
                                first = False
                            else:
                                run.text = run_text[:rs] + run_text[re:]
                        except Exception as e:
                            logger.warning(f"Erreur lors du remplacement dans le run: {e}")
                            continue
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement de l'entité {ent}: {e}")
                    continue
        except Exception as e:
            logger.error(f"Erreur lors du remplacement avec mapping: {e}")

    def anonymize_docx(self, data: bytes) -> Tuple[bytes, List[Entity], List[RunInfo], str]:
        """Anonymize a DOCX document while preserving structure and metadata.

        Returns the anonymized document bytes, detected entities, the mapping
        between plain text and DOCX runs and the plain text itself.
        """
        try:
            doc = Document(BytesIO(data))
            
            # Preserve metadata
            core_props = {}
            try:
                core_props = {
                    "author": doc.core_properties.author,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                }
            except Exception as e:
                logger.warning(f"Erreur lors de la sauvegarde des métadonnées: {e}")
            
            text, mapping = self.docx_text_mapping(doc)
            entities = self.detect(text)
            self._replace_using_mapping(doc, entities, mapping)

            # Restore metadata explicitly
            try:
                if core_props.get("author"):
                    doc.core_properties.author = core_props["author"]
                if core_props.get("title"):
                    doc.core_properties.title = core_props["title"]
                if core_props.get("subject"):
                    doc.core_properties.subject = core_props["subject"]
            except Exception as e:
                logger.warning(f"Erreur lors de la restauration des métadonnées: {e}")

            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read(), entities, mapping, text

        except Exception as e:
            logger.error(f"Erreur lors de l'anonymisation DOCX: {e}")
            raise

    def export_docx(
        self,
        data: bytes,
        mapping: Optional[List[RunInfo]] = None,
        entities: Optional[List[Entity]] = None,
        watermark: Optional[str] = None,
        audit: bool = False,
    ) -> Tuple[bytes, Optional[str]]:
        """Apply modifications and export options to a DOCX document."""
        try:
            doc = Document(BytesIO(data))

            if entities and mapping:
                self._replace_using_mapping(doc, entities, mapping)

            if watermark:
                try:
                    for section in doc.sections:
                        if hasattr(section, 'header') and section.header:
                            para = section.header.add_paragraph()
                            run = para.add_run(watermark)
                            run.font.color.rgb = None
                            run.font.bold = True
                except Exception as e:
                    logger.warning(f"Erreur lors de l'ajout du filigrane: {e}")

            report: Optional[str] = None
            if audit and entities:
                try:
                    report_lines = [f"{e.type}: {e.value}" for e in entities]
                    report = "\n".join(report_lines)
                except Exception as e:
                    logger.warning(f"Erreur lors de la création du rapport d'audit: {e}")

            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read(), report

        except Exception as e:
            logger.error(f"Erreur lors de l'export DOCX: {e}")
            raise

    # ------------------------------------------------------------------
    # PDF utilities
    def anonymize_pdf(
        self, data: bytes
    ) -> Tuple[bytes, List[Entity], List[RunInfo], str, bytes]:
        """Convert a PDF to DOCX, anonymize it and return mapping info.

        Additionally extracts bounding boxes for detected entities using
        ``pdfplumber`` so that the frontend can highlight them on the
        original PDF. Bounding boxes are expressed in the PDF coordinate
        system where the origin is at the top-left corner.
        """
        try:
            # Convert the PDF to DOCX and run the regular anonymization pipeline
            with tempfile.TemporaryDirectory() as tmpdir:
                pdf_path = Path(tmpdir) / "input.pdf"
                docx_path = Path(tmpdir) / "converted.docx"
                pdf_path.write_bytes(data)
                
                try:
                    pdf2docx_parse(str(pdf_path), str(docx_path))
                    original_docx = docx_path.read_bytes()
                except Exception as e:
                    logger.error(f"Erreur lors de la conversion PDF vers DOCX: {e}")
                    raise

            anonymized, entities, mapping, text = self.anonymize_docx(original_docx)

            # Compute bounding boxes from the original PDF. Any failure in this
            # auxiliary step should not prevent the main anonymization workflow.
            try:
                with pdfplumber.open(BytesIO(data)) as pdf:
                    char_map = []
                    pdf_text_parts: List[str] = []
                    pos = 0
                    for page_num, page in enumerate(pdf.pages, start=1):
                        try:
                            for ch in page.chars:
                                c = ch.get("text", "")
                                if not c:
                                    continue
                                pdf_text_parts.append(c)
                                char_map.append(
                                    {
                                        "index": pos,
                                        "page": page_num,
                                        "x0": ch.get("x0", 0),
                                        "x1": ch.get("x1", 0),
                                        "top": ch.get("top", 0),
                                        "bottom": ch.get("bottom", 0),
                                    }
                                )
                                pos += len(c)
                            pdf_text_parts.append("\n")
                            pos += 1
                        except Exception as e:
                            logger.warning(f"Erreur lors du traitement de la page {page_num}: {e}")
                            continue

                    pdf_text = "".join(pdf_text_parts)

                    search_pos = 0
                    for ent in sorted(entities, key=lambda e: e.start):
                        try:
                            idx = pdf_text.find(ent.value, search_pos)
                            if idx == -1:
                                continue
                            search_pos = idx + len(ent.value)
                            chars = [
                                cm for cm in char_map if idx <= cm["index"] < idx + len(ent.value)
                            ]
                            if not chars:
                                continue
                            x0 = min(c["x0"] for c in chars)
                            x1 = max(c["x1"] for c in chars)
                            top = min(c["top"] for c in chars)
                            bottom = max(c["bottom"] for c in chars)
                            ent.page = chars[0]["page"]
                            ent.x = x0
                            ent.y = top
                            ent.width = x1 - x0
                            ent.height = bottom - top
                        except Exception as e:
                            logger.warning(f"Erreur lors du calcul des coordonnées pour l'entité {ent.value}: {e}")
                            continue
            except Exception as e:
                # If anything goes wrong (e.g. malformed PDF), we simply skip
                # bounding box extraction and return the anonymized document.
                logger.warning(f"Erreur lors de l'extraction des bounding boxes: {e}")

            return anonymized, entities, mapping, text, original_docx

        except Exception as e:
            logger.error(f"Erreur lors de l'anonymisation PDF: {e}")
            raise   
    