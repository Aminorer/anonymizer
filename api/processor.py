import io
import logging
import os
from typing import BinaryIO, Tuple
from docx import Document
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    📄 Processeur de documents optimisé pour Vercel
    
    ✅ Support PDF (avec OCR) et DOCX
    ✅ Extraction de texte optimisée
    ✅ Conversion PDF→DOCX pour traitement uniforme
    ✅ Gestion d'erreurs robuste
    ✅ Limites Vercel respectées (<30s, <512MB)
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.ocr_dpi = 200  # Réduire pour Vercel (performance vs qualité)
        
        logger.info("📄 DocumentProcessor initialisé pour Vercel")
        
    def process_uploaded_file(self, file_content: bytes, filename: str) -> Tuple[Document, str]:
        """
        🔄 Traite un fichier uploadé et retourne un document DOCX et son texte
        
        Flux de traitement :
        PDF → (OCR si nécessaire) → Texte → DOCX
        DOCX → Extraction texte direct → DOCX de travail
        """
        # Validation de base
        if len(file_content) > self.max_file_size:
            raise ValueError(f"Fichier trop volumineux: {len(file_content) / (1024*1024):.1f}MB > 50MB")
        
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Format de fichier non supporté: {file_extension}")
        
        logger.info(f"📄 Traitement fichier: {filename} ({len(file_content) / (1024*1024):.1f}MB)")
        
        try:
            if file_extension == '.pdf':
                return self._process_pdf(file_content, filename)
            elif file_extension == '.docx':
                return self._process_docx(file_content, filename)
            else:
                raise ValueError(f"Extension non gérée: {file_extension}")
                
        except Exception as e:
            logger.error(f"❌ Erreur traitement {filename}: {e}")
            raise ValueError(f"Impossible de traiter le fichier: {str(e)}")
    
    def _process_pdf(self, pdf_content: bytes, filename: str) -> Tuple[Document, str]:
        """
        📖 Traite un fichier PDF avec extraction de texte et OCR si nécessaire
        """
        logger.info(f"📖 Traitement PDF: {filename}")
        
        try:
            # ÉTAPE 1 : Tentative d'extraction texte direct
            text = self._extract_text_from_pdf(pdf_content)
            logger.info(f"📝 Texte extrait directement: {len(text)} caractères")
            
            # ÉTAPE 2 : OCR si peu de texte extrait (PDF scanné)
            if len(text.strip()) < 100:
                logger.info("🔍 PDF scanné détecté, lancement OCR...")
                text = self._ocr_pdf_with_tesseract(pdf_content)
                logger.info(f"🔍 Texte OCR: {len(text)} caractères")
            
            # ÉTAPE 3 : Création DOCX à partir du texte
            docx_document = self._create_docx_from_text(text, filename)
            
            logger.info(f"✅ PDF traité avec succès: {len(text)} caractères")
            return docx_document, text
            
        except Exception as e:
            logger.error(f"❌ Erreur traitement PDF {filename}: {e}")
            raise ValueError(f"Impossible de traiter le PDF: {str(e)}")
    
    def _process_docx(self, docx_content: bytes, filename: str) -> Tuple[Document, str]:
        """
        📝 Traite un fichier DOCX avec extraction de texte
        """
        logger.info(f"📝 Traitement DOCX: {filename}")
        
        try:
            # Charger le document DOCX
            docx_stream = io.BytesIO(docx_content)
            document = Document(docx_stream)
            
            # Extraire le texte complet
            text = self._extract_text_from_docx(document)
            
            logger.info(f"✅ DOCX traité: {len(text)} caractères")
            return document, text
            
        except Exception as e:
            logger.error(f"❌ Erreur traitement DOCX {filename}: {e}")
            raise ValueError(f"Impossible de traiter le DOCX: {str(e)}")
    
    def _extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """
        📄 Extrait le texte d'un PDF de manière optimisée
        """
        text_parts = []
        
        try:
            pdf_stream = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            logger.info(f"📄 PDF: {len(pdf_reader.pages)} pages")
            
            # Limiter le nombre de pages pour Vercel (timeout)
            max_pages = min(len(pdf_reader.pages), 50)  # Limite pour performance
            
            for page_num, page in enumerate(pdf_reader.pages[:max_pages]):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                        
                    # Log progress pour longs documents
                    if (page_num + 1) % 10 == 0:
                        logger.info(f"📄 Traité {page_num + 1}/{max_pages} pages")
                        
                except Exception as e:
                    logger.warning(f"⚠️ Erreur page {page_num + 1}: {e}")
                    continue
            
            if max_pages < len(pdf_reader.pages):
                logger.warning(f"⚠️ PDF tronqué: {max_pages}/{len(pdf_reader.pages)} pages (limite Vercel)")
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"📄 Extraction PDF terminée: {len(extracted_text)} caractères")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"❌ Erreur extraction PDF: {e}")
            return ""
    
    def _ocr_pdf_with_tesseract(self, pdf_content: bytes) -> str:
        """
        🔍 Effectue l'OCR sur un PDF scanné (optimisé Vercel)
        """
        text_parts = []
        
        try:
            logger.info("🔍 Conversion PDF en images pour OCR...")
            
            # Conversion avec DPI réduit pour Vercel
            images = convert_from_bytes(
                pdf_content, 
                dpi=self.ocr_dpi,  # 200 DPI (compromis qualité/vitesse)
                first_page=1,
                last_page=20  # Limiter à 20 pages pour timeout Vercel
            )
            
            logger.info(f"🔍 {len(images)} images converties, lancement OCR...")
            
            for i, image in enumerate(images):
                try:
                    # OCR avec Tesseract français
                    page_text = pytesseract.image_to_string(
                        image, 
                        lang='fra',  # Français
                        config='--psm 6 --oem 1'  # Mode optimisé
                    )
                    
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                    
                    logger.info(f"🔍 OCR page {i+1}/{len(images)}: {len(page_text)} caractères")
                    
                except Exception as e:
                    logger.warning(f"⚠️ Erreur OCR page {i+1}: {e}")
                    continue
            
            ocr_text = "\n".join(text_parts)
            logger.info(f"✅ OCR terminé: {len(ocr_text)} caractères")
            
            return ocr_text
            
        except Exception as e:
            logger.error(f"❌ Erreur OCR: {e}")
            return ""
    
    def _extract_text_from_docx(self, document: Document) -> str:
        """
        📝 Extrait le texte complet d'un document DOCX (paragraphes + tableaux)
        """
        text_parts = []
        
        try:
            # PARTIE 1 : Extraire le texte des paragraphes
            paragraph_count = 0
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text_parts.append(paragraph_text)
                    paragraph_count += 1
            
            logger.debug(f"📝 {paragraph_count} paragraphes extraits")
            
            # PARTIE 2 : Extraire le texte des tableaux
            table_count = 0
            for table in document.tables:
                table_texts = []
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        table_texts.append(" | ".join(row_texts))
                
                if table_texts:
                    text_parts.append("\n".join(table_texts))
                    table_count += 1
            
            logger.debug(f"📝 {table_count} tableaux extraits")
            
            # PARTIE 3 : Assembler le texte final
            full_text = "\n\n".join(text_parts)
            
            logger.info(f"📝 Extraction DOCX terminée: {len(full_text)} caractères")
            return full_text
            
        except Exception as e:
            logger.error(f"❌ Erreur extraction texte DOCX: {e}")
            return ""
    
    def _create_docx_from_text(self, text: str, original_filename: str) -> Document:
        """
        📄 Crée un document DOCX à partir d'un texte (pour PDFs convertis)
        """
        try:
            document = Document()
            
            # Ajouter un en-tête avec le nom du fichier original
            title = document.add_heading(f'Document converti: {original_filename}', 0)
            
            # Ajouter une note sur la conversion
            info_paragraph = document.add_paragraph()
            info_paragraph.add_run("Document converti automatiquement pour anonymisation. ").italic = True
            info_paragraph.add_run("Format original préservé lors de l'export final.").italic = True
            
            # Diviser le texte en paragraphes intelligemment
            paragraphs = self._smart_split_text(text)
            
            logger.info(f"📄 Création DOCX: {len(paragraphs)} paragraphes")
            
            # Ajouter chaque paragraphe
            for i, paragraph_text in enumerate(paragraphs):
                if paragraph_text.strip():
                    paragraph = document.add_paragraph(paragraph_text.strip())
                    
                    # Progress pour longs textes
                    if (i + 1) % 100 == 0:
                        logger.debug(f"📄 Ajouté {i + 1}/{len(paragraphs)} paragraphes")
            
            logger.info(f"✅ Document DOCX créé: {len(paragraphs)} paragraphes")
            return document
            
        except Exception as e:
            logger.error(f"❌ Erreur création DOCX: {e}")
            raise ValueError(f"Impossible de créer le document DOCX: {str(e)}")
    
    def _smart_split_text(self, text: str) -> list[str]:
        """
        🧠 Découpe intelligente du texte en paragraphes
        """
        # Normaliser les sauts de ligne
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Séparer par double saut de ligne d'abord
        paragraphs = text.split('\n\n')
        
        # Affiner la séparation
        refined_paragraphs = []
        for paragraph in paragraphs:
            # Si le paragraphe est très long, le subdiviser
            if len(paragraph) > 2000:
                # Séparer par phrases
                sentences = paragraph.split('. ')
                current_para = ""
                
                for sentence in sentences:
                    if len(current_para) + len(sentence) > 1000:
                        if current_para:
                            refined_paragraphs.append(current_para.strip())
                        current_para = sentence + ". "
                    else:
                        current_para += sentence + ". "
                
                if current_para:
                    refined_paragraphs.append(current_para.strip())
            else:
                if paragraph.strip():
                    refined_paragraphs.append(paragraph.strip())
        
        return refined_paragraphs
    
    def apply_global_replacements(self, document: Document, replacements: dict) -> bytes:
        """
        🔄 Applique les remplacements globaux dans un document DOCX
        
        Préserve la mise en forme originale tout en anonymisant le contenu
        """
        try:
            logger.info(f"🔄 Application de {len(replacements)} remplacements")
            
            replacement_count = 0
            
            # PARTIE 1 : Remplacements dans les paragraphes
            for paragraph in document.paragraphs:
                original_text = paragraph.text
                
                for original, replacement in replacements.items():
                    if original in paragraph.text:
                        # Remplacer dans chaque run pour préserver la mise en forme
                        for run in paragraph.runs:
                            if original in run.text:
                                run.text = run.text.replace(original, replacement)
                                replacement_count += 1
            
            # PARTIE 2 : Remplacements dans les tableaux
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for original, replacement in replacements.items():
                                if original in paragraph.text:
                                    for run in paragraph.runs:
                                        if original in run.text:
                                            run.text = run.text.replace(original, replacement)
                                            replacement_count += 1
            
            # PARTIE 3 : Sauvegarder en mémoire
            output_stream = io.BytesIO()
            document.save(output_stream)
            output_stream.seek(0)
            
            document_bytes = output_stream.read()
            output_stream.close()
            
            logger.info(f"✅ {replacement_count} remplacements appliqués, document généré ({len(document_bytes) / (1024*1024):.1f}MB)")
            
            return document_bytes
            
        except Exception as e:
            logger.error(f"❌ Erreur application remplacements: {e}")
            raise ValueError(f"Impossible d'appliquer les remplacements: {str(e)}")

# Instance globale pour l'application
document_processor = DocumentProcessor()